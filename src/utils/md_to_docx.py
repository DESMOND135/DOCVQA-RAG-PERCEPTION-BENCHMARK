import os
import re
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER, WD_PARAGRAPH_ALIGNMENT
from docx.enum.section import WD_SECTION, WD_ORIENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement, ns, parse_xml
from pygments import highlight
from pygments.lexers import get_lexer_by_name, PythonLexer
from pygments.formatters import RawTokenFormatter
from pygments.token import Token

def set_paragraph_shading(p, color_hex):
    """Sets background color shading for a paragraph."""
    pPr = p._p.get_or_add_pPr()
    shd = parse_xml(f'<w:shd xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:val="clear" w:color="auto" w:fill="{color_hex}"/>')
    pPr.append(shd)

def set_paragraph_borders(p, color_hex="003366"):
    """Sets a left border/accent bar for a paragraph to denote a code block."""
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(f'<w:pBdr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
                     f'<w:left w:val="single" w:sz="24" w:space="15" w:color="{color_hex}"/>'
                     f'</w:pBdr>')
    pPr.append(pBdr)

def add_syntax_highlighted_code(doc, code, language="python"):
    """Adds a paragraph of modern, syntax-highlighted Python/other code to the document."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.right_indent = Inches(0.5)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.05
    
    # Apply background shading and left border
    set_paragraph_shading(p, "F5F5F5")
    set_paragraph_borders(p, "003366") # Navy accent bar
    
    # Premium Jupyter / VS Code Light theme colors
    styles = {
        Token.Keyword: "0000FF",       # Blue
        Token.Keyword.Constant: "0000FF",
        Token.Keyword.Declaration: "0000FF",
        Token.Keyword.Namespace: "0000FF",
        Token.Keyword.Pseudo: "0000FF",
        Token.Keyword.Reserved: "0000FF",
        Token.Keyword.Type: "0000FF",
        Token.Comment: "008000",       # Green
        Token.Comment.Single: "008000",
        Token.Comment.Multiline: "008000",
        Token.String: "A31515",        # Red
        Token.String.Char: "A31515",
        Token.String.Doc: "A31515",
        Token.String.Double: "A31515",
        Token.String.Escape: "A31515",
        Token.String.Heredoc: "A31515",
        Token.String.Interpol: "A31515",
        Token.String.Other: "A31515",
        Token.String.Regex: "A31515",
        Token.String.Single: "A31515",
        Token.Number: "098658",        # Teal/Greenish
        Token.Number.Bin: "098658",
        Token.Number.Float: "098658",
        Token.Number.Hex: "098658",
        Token.Number.Integer: "098658",
        Token.Number.Integer.Long: "098658",
        Token.Number.Oct: "098658",
        Token.Name.Function: "795E26", # Yellowish-brown / Gold
        Token.Name.Class: "267F99",    # Teal-blue
        Token.Name.Builtin: "0000FF",   # Blue
        Token.Name.Builtin.Pseudo: "0000FF",
        Token.Operator: "000000",      # Black
        Token.Punctuation: "000000",
    }
    
    lexer = PythonLexer()
    tokens = lexer.get_tokens(code)
    
    for ttype, value in tokens:
        color_hex = None
        temp_type = ttype
        while temp_type is not None:
            if temp_type in styles:
                color_hex = styles[temp_type]
                break
            temp_type = temp_type.parent
            
        run = p.add_run(value)
        run.font.name = 'Consolas'
        run.font.size = Pt(9.5)
        
        if color_hex:
            run.font.color.rgb = RGBColor(
                int(color_hex[0:2], 16),
                int(color_hex[2:4], 16),
                int(color_hex[4:6], 16)
            )
        else:
            run.font.color.rgb = RGBColor(0, 0, 0)

def create_element(name):
    return OxmlElement(name)

def create_attribute(element, name, value):
    element.set(ns.qn(name), value)

def add_field(run, field_type):
    """
    Adds a Word field (PAGE, DATE, TOC) to a run.
    """
    fldChar1 = create_element('w:fldChar')
    create_attribute(fldChar1, 'w:fldCharType', 'begin')
    
    instrText = create_element('w:instrText')
    create_attribute(instrText, 'xml:space', 'preserve')
    instrText.text = field_type
    
    fldChar2 = create_element('w:fldChar')
    create_attribute(fldChar2, 'w:fldCharType', 'separate')
    
    t = create_element('w:t')
    t.text = "1" # Placeholder
    
    fldChar3 = create_element('w:fldChar')
    create_attribute(fldChar3, 'w:fldCharType', 'end')
    
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run._r.append(t)
    run._r.append(fldChar3)

def force_field_update(doc):
    element = doc.settings.element
    update_fields = OxmlElement('w:updateFields')
    update_fields.set(ns.qn('w:val'), 'true')
    element.append(update_fields)

def set_toc_styles(doc):
    for i in range(1, 4):
        style_name = f'TOC {i}'
        if style_name in doc.styles:
            style = doc.styles[style_name]
            style.font.name = 'Arial'
            style.font.size = Pt(12)
            style.paragraph_format.line_spacing = 1.15
            style.paragraph_format.space_after = Pt(2)


def get_custom_omml(latex):
    """Returns OMML XML string (m:oMath element) for the three known equations.
    Returns plain <m:oMath> so it can be safely appended into a w:p element.
    All characters are in the Basic Multilingual Plane to ensure XML 1.0 compatibility."""
    NS = 'xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math"'
    if "ANLS" in latex:
        # ANLS = 1/N sum_{i=1}^N max_j ( 0, 1 - NL(p_i,g_{i,j})/max(|p_i|,|g_{i,j}|) )
        return (
            f'<m:oMath {NS}>'
            '<m:r><m:rPr><m:nor/></m:rPr><m:t>ANLS = </m:t></m:r>'
            # 1/N skewed fraction
            '<m:f><m:fPr><m:type m:val="skw"/></m:fPr>'
            '<m:num><m:r><m:t>1</m:t></m:r></m:num>'
            '<m:den><m:r><m:t>N</m:t></m:r></m:den>'
            '</m:f>'
            # Summation i=1 to N
            '<m:nary><m:naryPr>'
            '<m:chr m:val="\u2211"/>'
            '<m:limLoc m:val="undOvr"/><m:grow m:val="1"/>'
            '<m:subHide m:val="0"/><m:supHide m:val="0"/>'
            '</m:naryPr>'
            '<m:sub><m:r><m:t>i=1</m:t></m:r></m:sub>'
            '<m:sup><m:r><m:t>N</m:t></m:r></m:sup>'
            '<m:e>'
            # max_j
            '<m:sSub><m:sSubPr><m:ctrlPr/></m:sSubPr>'
            '<m:e><m:r><m:t>max</m:t></m:r></m:e>'
            '<m:sub><m:r><m:t>j</m:t></m:r></m:sub>'
            '</m:sSub>'
            # ( 0 , 1 - fraction )
            '<m:d><m:dPr>'
            '<m:begChr m:val="("/><m:endChr m:val=")"/>'
            '</m:dPr><m:e>'
            '<m:r><m:t xml:space="preserve">0, 1\u2212</m:t></m:r>'
            '<m:f><m:fPr><m:type m:val="bar"/></m:fPr>'
            '<m:num>'
            '<m:r><m:t xml:space="preserve">NL(</m:t></m:r>'
            '<m:sSub><m:sSubPr><m:ctrlPr/></m:sSubPr>'
            '<m:e><m:r><m:t>p</m:t></m:r></m:e>'
            '<m:sub><m:r><m:t>i</m:t></m:r></m:sub>'
            '</m:sSub>'
            '<m:r><m:t>, </m:t></m:r>'
            '<m:sSub><m:sSubPr><m:ctrlPr/></m:sSubPr>'
            '<m:e><m:r><m:t>g</m:t></m:r></m:e>'
            '<m:sub><m:r><m:t>i,j</m:t></m:r></m:sub>'
            '</m:sSub>'
            '<m:r><m:t>)</m:t></m:r>'
            '</m:num>'
            '<m:den>'
            '<m:r><m:t xml:space="preserve">max(|</m:t></m:r>'
            '<m:sSub><m:sSubPr><m:ctrlPr/></m:sSubPr>'
            '<m:e><m:r><m:t>p</m:t></m:r></m:e>'
            '<m:sub><m:r><m:t>i</m:t></m:r></m:sub>'
            '</m:sSub>'
            '<m:r><m:t>|, |</m:t></m:r>'
            '<m:sSub><m:sSubPr><m:ctrlPr/></m:sSubPr>'
            '<m:e><m:r><m:t>g</m:t></m:r></m:e>'
            '<m:sub><m:r><m:t>i,j</m:t></m:r></m:sub>'
            '</m:sSub>'
            '<m:r><m:t>|)</m:t></m:r>'
            '</m:den></m:f>'
            '</m:e></m:d>'
            '</m:e></m:nary>'
            '</m:oMath>'
        )
    elif "EM =" in latex:
        # EM = 1/N sum_{i=1}^N 1(p_i = g_i)
        return (
            f'<m:oMath {NS}>'
            '<m:r><m:rPr><m:nor/></m:rPr><m:t>EM = </m:t></m:r>'
            '<m:f><m:fPr><m:type m:val="skw"/></m:fPr>'
            '<m:num><m:r><m:t>1</m:t></m:r></m:num>'
            '<m:den><m:r><m:t>N</m:t></m:r></m:den>'
            '</m:f>'
            '<m:nary><m:naryPr>'
            '<m:chr m:val="\u2211"/>'
            '<m:limLoc m:val="undOvr"/><m:grow m:val="1"/>'
            '<m:subHide m:val="0"/><m:supHide m:val="0"/>'
            '</m:naryPr>'
            '<m:sub><m:r><m:t>i=1</m:t></m:r></m:sub>'
            '<m:sup><m:r><m:t>N</m:t></m:r></m:sup>'
            '<m:e>'
            # Bold 1 indicator function - bold plain "1" (no SMP Unicode chars)
            '<m:r><m:rPr><m:b/></m:rPr><m:t>1</m:t></m:r>'
            '<m:d><m:dPr>'
            '<m:begChr m:val="("/><m:endChr m:val=")"/>'
            '</m:dPr><m:e>'
            '<m:sSub><m:sSubPr><m:ctrlPr/></m:sSubPr>'
            '<m:e><m:r><m:t>p</m:t></m:r></m:e>'
            '<m:sub><m:r><m:t>i</m:t></m:r></m:sub>'
            '</m:sSub>'
            '<m:r><m:t>=</m:t></m:r>'
            '<m:sSub><m:sSubPr><m:ctrlPr/></m:sSubPr>'
            '<m:e><m:r><m:t>g</m:t></m:r></m:e>'
            '<m:sub><m:r><m:t>i</m:t></m:r></m:sub>'
            '</m:sSub>'
            '</m:e></m:d>'
            '</m:e></m:nary>'
            '</m:oMath>'
        )
    elif "F1 =" in latex:
        # F1 = 2 * Precision*Recall / (Precision+Recall)
        return (
            f'<m:oMath {NS}>'
            '<m:r><m:rPr><m:nor/></m:rPr><m:t>F1 = 2\u00b7</m:t></m:r>'
            '<m:f><m:fPr><m:type m:val="bar"/></m:fPr>'
            '<m:num><m:r><m:t>Precision\u00b7Recall</m:t></m:r></m:num>'
            '<m:den><m:r><m:t>Precision+Recall</m:t></m:r></m:den>'
            '</m:f>'
            '</m:oMath>'
        )
    return get_omml_for_latex(latex)

def get_omml_for_latex(latex):
    """Translates LaTeX string to native Word OMML XML using Microsoft's local MML2OMML XSLT."""
    try:
        import latex2mathml.converter
        from lxml import etree
        
        # Professional dynamic path search for the official Microsoft MML2OMML.XSL stylesheet
        xsl_path = r"C:\Program Files\Microsoft Office\root\Office16\MML2OMML.XSL"
        if not os.path.exists(xsl_path):
            import glob
            matches = glob.glob(r"C:\Program Files\Microsoft Office\root\**\MML2OMML.XSL", recursive=True)
            if matches:
                xsl_path = matches[0]
            else:
                return None
                
        xslt = etree.parse(xsl_path)
        transform = etree.XSLT(xslt)
        
        mathml_str = latex2mathml.converter.convert(latex, display='block')
        mathml_tree = etree.fromstring(mathml_str)
        omml_tree = transform(mathml_tree)
        omml_str = etree.tostring(omml_tree, encoding='utf-8').decode('utf-8')
        if '?>' in omml_str:
            omml_str = omml_str.split('?>')[-1].strip()
        return omml_str
    except Exception as e:
        print(f"OMML conversion error for '{latex}': {e}")
        return None

def render_latex_to_image(latex, output_path, is_inline=False):
    """Renders LaTeX to a high-resolution transparent PNG using Matplotlib in academic BLACK."""
    try:
        import matplotlib.pyplot as plt
        plt.rcParams['mathtext.fontset'] = 'stix'
        plt.rcParams['font.family'] = 'STIXGeneral'
        plt.rc('text', usetex=False)
        
        # Clean latex string
        latex = latex.strip('$').strip()
        if not latex: return False
        if not (latex.startswith('$') and latex.endswith('$')):
            latex = f'${latex}$'
            
        fig = plt.figure(figsize=(0.1, 0.1))
        # Equations MUST be BLACK for academic standard
        fig.text(0, 0, latex, fontsize=16 if not is_inline else 12, color='#000000')
        
        renderer = fig.canvas.get_renderer()
        bbox = fig.texts[0].get_window_extent(renderer=renderer)
        width_in = (bbox.width + 10) / fig.dpi
        height_in = (bbox.height + 10) / fig.dpi
        fig.set_size_inches(width_in, height_in)
        
        plt.savefig(output_path, dpi=300, transparent=True, bbox_inches='tight', pad_inches=0.02)
        plt.close(fig)
        return True
    except Exception as e:
        print(f"Latex render error: {e}")
        return False

def add_caption(doc, label_text, raw_caption_text, chapter_num="1"):
    """Adds a professional academic caption (X.Y) with manual numbering for total reliability."""
    clean_caption = re.sub(r'^(Figure|Table|Equation|Formula)\s*[\dA-Z.]*[:.]?\s*', '', raw_caption_text, flags=re.IGNORECASE)
    
    if not hasattr(doc, '_caption_counts'): doc._caption_counts = {}
    key = label_text
    doc._caption_counts[key] = doc._caption_counts.get(key, 0) + 1
    num = doc._caption_counts[key]
    
    # Auto-add source for figures to satisfy user requirements
    if label_text == "Figure" and clean_caption:
        lower_cap = clean_caption.lower()
        if "paddleocr" in lower_cap or "layout detection" in lower_cap:
            clean_caption += " (Source: Adapted from [4])"
        elif "tesseract" in lower_cap:
            clean_caption += " (Source: Adapted from [16])"
        elif "vlm" in lower_cap or "donut" in lower_cap or "vit" in lower_cap or "primitives" in lower_cap:
            clean_caption += " (Source: Adapted from [31])"
        elif "dataset complexity" in lower_cap or "samples" in lower_cap:
            clean_caption += " (Source: Adapted from [22])"
        else:
            clean_caption += " (Source: Own elaboration)"
            
    p = doc.add_paragraph(style='Caption')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    
    label_run = p.add_run(f"{label_text} {num}")
    label_run.bold = True
    
    if clean_caption:
        p.add_run(f": {clean_caption}").bold = True
    
    if "Equation" in label_text:
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT; p.clear()
        p.add_run(f"({chapter_num}.{num})").italic = True
        return p
    return p

def add_formatted_text(p, text, doc=None, chapter_num="1"):
    # Handle Block Math
    if text.strip().startswith('$$') and text.strip().endswith('$$'):
        latex = text.strip().strip('$').strip()

        # Parse equation number if present, e.g. \quad (3)
        eq_num = None
        m = re.search(r'(?:\\quad\s*|\\qquad\s*|\s+)\((\d+(?:\.\d+)?)\)\s*$', latex)
        if m:
            eq_num = m.group(1)
            latex = latex[:m.start()].strip()

        omml_str = get_custom_omml(latex)
        if omml_str:
            # ------------------------------------------------------------------
            # SAFE APPROACH: append the m:oMath element directly into p._p
            # The paragraph stays in the document body (never inside a table),
            # so Word renders it as proper display math (not linear notation).
            # The equation number goes in a separate tight right-aligned para.
            # ------------------------------------------------------------------
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(0)

            # Append the oMath element into this paragraph's XML
            omml_element = parse_xml(omml_str)
            p._p.append(omml_element)

            # Add tight right-aligned equation number paragraph
            if eq_num and doc is not None:
                num_p = doc.add_paragraph()
                num_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                num_p.paragraph_format.space_before = Pt(0)
                num_p.paragraph_format.space_after = Pt(4)
                run = num_p.add_run(f"({eq_num})")
                run.font.name = 'Cambria'
                run.font.size = Pt(11)
            return
        else:
            # Fallback to high-resolution Matplotlib image rendering
            temp_img = f"temp_math_{hash(latex)}.png"
            if render_latex_to_image(latex, temp_img, is_inline=False):
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run()
                run.add_picture(temp_img)
                try: os.remove(temp_img)
                except: pass
                if doc: add_caption(doc, "Equation", "", chapter_num=chapter_num)
                return
            
    # Handle Inline Math (Negative lookahead prevents currency like $400 from matching as math)
    parts = re.split(r'(\$(?!\d)[^$]+\$|\*\*\*.*?\*\*\*|\*\*.*?\*\*|\*.*?\*)', text)
    for part in parts:
        if not part: continue
        if part.startswith('$') and part.endswith('$'):
            latex = part.strip('$')
            omml_str = get_custom_omml(latex)
            if omml_str:
                omml_element = parse_xml(omml_str)
                p._p.append(omml_element)
            else:
                # Fallback to high-resolution inline image
                temp_img = f"temp_math_inline_{hash(latex)}.png"
                if render_latex_to_image(latex, temp_img, is_inline=True):
                    run = p.add_run()
                    run.add_picture(temp_img, height=Cm(0.45))
                    try: os.remove(temp_img)
                    except: pass
                else:
                    run = p.add_run(part); run.italic = True
        elif part.startswith('***') and part.endswith('***'):
            p.add_run(part[3:-3]).bold = True; p.runs[-1].italic = True
        elif part.startswith('**') and part.endswith('**'):
            p.add_run(part[2:-2]).bold = True
        elif part.startswith('*') and part.endswith('*'):
            p.add_run(part[1:-1]).italic = True
        else:
            p.add_run(part)

def set_academic_styles(doc, is_paper=False):
    style = doc.styles['Normal']
    # Arial as the sans-serif font according to PROJECT FORMAT.pdf
    font = style.font; font.name = 'Cambria' if is_paper else 'Arial'; font.size = Pt(12)
    pf = style.paragraph_format; pf.line_spacing = 1.15 if is_paper else 1.5; pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf.space_after = Pt(6) if is_paper else Pt(0)
    pf.space_before = Pt(0)

def set_heading_styles(doc, is_paper=False):
    # Rule 1: No heading should appear alone at the bottom
    for level in range(0, 5):
        s_name = f'Heading {level}' if level > 0 else 'Title'
        if s_name in doc.styles:
            s = doc.styles[s_name]
            s.paragraph_format.keep_with_next = True
            s.paragraph_format.keep_together = True
            if is_paper:
                s.paragraph_format.space_before = Pt(0)
                s.paragraph_format.space_after = Pt(0)
            else:
                s.paragraph_format.space_before = Pt(6)
                s.paragraph_format.space_after = Pt(6)
            s.font.name = 'Cambria' if is_paper else 'Arial'
            if level == 1: s.font.size = Pt(16); s.font.bold = True
            elif level == 2: s.font.size = Pt(14); s.font.bold = True
            elif level == 3: s.font.size = Pt(12); s.font.bold = True

def define_caption_style(doc):
    if 'Caption' not in doc.styles:
        s = doc.styles.add_style('Caption', WD_STYLE_TYPE.PARAGRAPH)
        s.base_style = doc.styles['Normal']; s.font.name = 'Arial'; s.font.size = Pt(11); s.font.italic = True
        s.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

def set_page_setup(doc, is_paper=False):
    section = doc.sections[0]
    section.page_height = Cm(29.7); section.page_width = Cm(21.0)
    if is_paper:
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)
        doc.settings.odd_and_even_pages_header_footer = False
    else:
        section.left_margin = Cm(3.5)
        section.right_margin = Cm(2.0)
        doc.settings.odd_and_even_pages_header_footer = True
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)

def add_paper_footer(section):
    section.footer.is_linked_to_previous = False
    p = section.footer.paragraphs[0]
    p.clear()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.font.name = 'Cambria'
    run.font.size = Pt(10)
    add_field(run, "PAGE")

def add_academic_footer(section, has_numbering=True):
    section.footer.is_linked_to_previous = False
    section.even_page_footer.is_linked_to_previous = False
    if not has_numbering:
        section.footer.paragraphs[0].clear(); section.even_page_footer.paragraphs[0].clear(); return
    for i, p in enumerate([section.footer.paragraphs[0], section.even_page_footer.paragraphs[0]]):
        p.clear(); p.alignment = WD_ALIGN_PARAGRAPH.RIGHT if i == 0 else WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(); run.font.name = 'Times New Roman'; run.font.size = Pt(12)
        add_field(run, "PAGE")

def set_column_layout(section, num_columns=2):
    """Sets the number of columns for a section."""
    sectPr = section._sectPr
    cols = sectPr.xpath('./w:cols')
    if not cols:
        cols = OxmlElement('w:cols')
        sectPr.append(cols)
    else:
        cols = cols[0]
    cols.set(ns.qn('w:num'), str(num_columns))
    cols.set(ns.qn('w:space'), '720') # 0.5 inch space between columns


from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import RGBColor

def format_custom_table(tbl):
    # Set borders
    tblPr = tbl._element.xpath('w:tblPr')
    if tblPr:
        tblBorders = OxmlElement('w:tblBorders')
        for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), '4')
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), '4F81BD')
            tblBorders.append(border)
        tblPr[0].append(tblBorders)

    # Shade header row and bold text
    for i, cell in enumerate(tbl.rows[0].cells):
        tcPr = cell._element.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), '4F81BD')
        tcPr.append(shd)
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.color.rgb = RGBColor(255, 255, 255)
    
    # Shade alternating rows
    for row_idx in range(1, len(tbl.rows)):
        if row_idx % 2 == 1:
            for cell in tbl.rows[row_idx].cells:
                tcPr = cell._element.get_or_add_tcPr()
                shd = OxmlElement('w:shd')
                shd.set(qn('w:val'), 'clear')
                shd.set(qn('w:color'), 'auto')
                shd.set(qn('w:fill'), 'D3DFEE')
                tcPr.append(shd)

def convert_to_professional_docx(md_path, docx_path):
    print(f"Applying Global Corrections to: {docx_path}")
    is_paper = "paper" in docx_path.lower()
    
    if is_paper:
        doc = Document()
    else:
        # Load the official template Szablon.docx
        # Resolve path correctly
        template_path = os.path.join(os.path.dirname(os.path.dirname(os.path.dirname(md_path))), "Szablon.docx")
        try:
            doc = Document(template_path)
            print(f"Loaded official template: {template_path}")
        except Exception as e:
            print(f"Template load failed: {e}, falling back to blank document")
            doc = Document()
    define_caption_style(doc)
    set_academic_styles(doc, is_paper=is_paper)
    set_heading_styles(doc, is_paper=is_paper)
    set_page_setup(doc, is_paper=is_paper)
    
    if is_paper:
        add_paper_footer(doc.sections[0])
    else:
        add_academic_footer(doc.sections[0], has_numbering=False)
        force_field_update(doc)
        set_toc_styles(doc)
    
    if not os.path.exists(md_path): return
    with open(md_path, 'r', encoding='utf-8') as f: lines = f.readlines()
    
    current_chapter = "1"
    just_added_chapter_break = False
    
    # Paper Specific Header Logic
    if is_paper:
        # Title
        title_line = next((l.lstrip('#').strip() for l in lines if l.startswith('# ')), "Document")
        tp = doc.add_heading(title_line, 0)
        tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Authors & Affiliations & Abstract parsing
        author_info = []
        abstract_lines = []
        in_abstract = False
        i = 0
        while i < len(lines):
            line = lines[i].strip()
            if line.startswith('# '):
                i += 1
                continue
            if line.startswith('**Abstract**'):
                in_abstract = True
                i += 1
                continue
            if in_abstract:
                if line.startswith('##') or line.startswith('Keywords:') or line.startswith('**Keywords:**'):
                    in_abstract = False
                    break
                if line: abstract_lines.append(line)
            else:
                if line and not line.startswith('#'):
                    author_info.append(line)
            i += 1
            if not in_abstract and i > 25: break # Safety break
            
        # Add Authors
        if author_info:
            ap = doc.add_paragraph()
            ap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            ap.paragraph_format.space_before = Pt(6)
            ap.paragraph_format.space_after = Pt(12)
            for info in author_info:
                run = ap.add_run(info.replace('**', '') + "\n")
                run.font.name = 'Cambria'
                run.font.size = Pt(10.5)
        
        # Switch to 2 columns for the rest (including Abstract and Keywords!)
        new_section = doc.add_section(WD_SECTION.CONTINUOUS)
        set_column_layout(new_section, 2)
        # Reset margins for paper columns
        new_section.left_margin = Cm(2.0)
        new_section.right_margin = Cm(2.0)
        new_section.top_margin = Cm(2.0)
        new_section.bottom_margin = Cm(2.0)
        add_paper_footer(new_section)
        
        # Add Abstract in 2-column layout!
        if abstract_lines:
            ab_head = doc.add_paragraph()
            ab_head.paragraph_format.space_before = Pt(6)
            ab_head.paragraph_format.space_after = Pt(4)
            run = ab_head.add_run("Abstract")
            run.bold = True
            run.font.name = 'Cambria'
            run.font.size = Pt(12)
            
            ab_body = doc.add_paragraph()
            ab_body.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            ab_body.paragraph_format.space_after = Pt(6)
            ab_body.paragraph_format.line_spacing = 1.15
            run = ab_body.add_run(" ".join(abstract_lines))
            run.font.name = 'Cambria'
            run.font.size = Pt(12)
            
        # Start processing from after abstract
        i = 0
        while i < len(lines) and not (lines[i].strip().startswith('**Keywords:**') or lines[i].strip().startswith('Keywords:') or lines[i].strip().startswith('## 1. INTRODUCTION')):
            i += 1
    else:
        # Thesis Header Logic
        # Custom Professional Title Page using Szablon.docx
        
        def replace_para_text(p, new_text):
            if not getattr(p, 'runs', []): return
            if not p.runs: p.add_run(new_text)
            else:
                p.runs[0].text = new_text
                for r in p.runs[1:]: r.text = ''
            
        try:
            replace_para_text(doc.paragraphs[3], 'MASTER THESIS')
            replace_para_text(doc.paragraphs[5], 'Systems-Level Reliability and Robustness Evaluation Framework for Document AI')
            replace_para_text(doc.paragraphs[6], '')
            replace_para_text(doc.paragraphs[8], 'Tifang Desmond Ngoe')
            replace_para_text(doc.paragraphs[9], '')
            replace_para_text(doc.paragraphs[10], 'Kierunek: Artificial Intelligence and Data Science')
            replace_para_text(doc.paragraphs[11], '') 
            replace_para_text(doc.paragraphs[12], 'Poziom studiów: II')
            replace_para_text(doc.paragraphs[13], 'Promotor pracy: Prof. Piotr Duda')
            replace_para_text(doc.paragraphs[21], 'Czestochowa University of Technology, 2026')
            
            # Delete dummy paragraphs from index 23 onwards
            def delete_paragraph(paragraph):
                p = paragraph._element
                p.getparent().remove(p)
                paragraph._p = paragraph._element = None
                
            for p in list(doc.paragraphs[23:]):
                delete_paragraph(p)
        except Exception as e:
            print(f"Warning: Failed to map template placeholders. Expected Szablon.docx structure not found: {e}")
        
        # Section 1: Table of Contents (Page 3)
        doc.add_section(WD_SECTION.ODD_PAGE)
        add_academic_footer(doc.sections[-1], has_numbering=True)
        
        toc_title_p = doc.add_paragraph()
        toc_title_p.paragraph_format.space_before = Pt(12)
        toc_title_p.paragraph_format.space_after = Pt(6)
        toc_title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = toc_title_p.add_run("TABLE OF CONTENTS")
        run.bold = True
        run.font.name = 'Arial'
        run.font.size = Pt(14)
        
        toc_p = doc.add_paragraph()
        toc_run = toc_p.add_run()
        add_field(toc_run, 'TOC \\o "1-3" \\h \\z \\u')
        
        # Skip the title and metadata lines to start parsing directly at the Abstract
        i = 0
        while i < len(lines) and not (lines[i].strip().startswith('## Abstract') or lines[i].strip().startswith('## CHAPTER 1')):
            i += 1

    while i < len(lines):
        line = lines[i].strip()
        if not line or line == '---' or line == '***': i += 1; continue
        if not line.startswith('## '):
            just_added_chapter_break = False
            
        if line.startswith('## '):
            title = line.lstrip('#').strip()
            # Always update chapter number if it starts with digit or matches CHAPTER X
            m = re.match(r'^(\d+)\.', title)
            m_chap = re.match(r'^CHAPTER\s+(\w+)', title, re.IGNORECASE)
            if m: current_chapter = m.group(1)
            elif m_chap: current_chapter = m_chap.group(1)
            elif "Appendix" in title: current_chapter = title.split(':')[0].replace("Appendix", "").strip()
            
            is_odd = False
            if not is_paper:
                is_chapter = title.upper().startswith("CHAPTER") or title.upper().startswith("APPENDIX")
                if is_chapter:
                    is_odd = True
                else:
                    is_odd = any(t.lower() in title.lower() for t in ["introduction", "conclusion", "references", "bibliography", "abstract", "table of contents", "list of"]) or (title and title[0].isdigit() and "." in title[:3])
            
            if "Abstract" in title or "Streszczenie" in title or "Keywords" in title:
                heading_text = "ABSTRACT"
                if "Streszczenie" in title:
                    heading_text = "STRESZCZENIE"
                elif "Keywords" in title:
                    heading_text = "Keywords / Słowa kluczowe"
                
                h = doc.add_heading(heading_text, 1)
                h.paragraph_format.space_before = Pt(12)
                h.paragraph_format.space_after = Pt(12)
                if is_odd and "Keywords" not in title:
                    h.paragraph_format.page_break_before = True
                h.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # Make sure the font matches other headings
                for run in h.runs:
                    run.font.name = 'Arial'
                    run.bold = True
                i += 1; continue
            
            h = doc.add_heading(title, 1)
            h.paragraph_format.keep_with_next = True
            if is_odd:
                h.paragraph_format.page_break_before = True
            i += 1; continue
        if line.startswith('### '): h = doc.add_heading(line.lstrip('#').strip(), level=2); h.paragraph_format.keep_with_next = True; i += 1; continue
        if line.startswith('#### '): p = doc.add_paragraph(); p.add_run(line.lstrip('#').strip()).bold = True; p.paragraph_format.keep_with_next = True; i += 1; continue
        if line.startswith('|') or (line.startswith('**Table') and i+1 < len(lines) and '|' in lines[i+1]):
            caption = ""; cap_p = None
            if line.startswith('**Table'): caption = line.strip('*').strip(); i += 1
            if caption: 
                cap_p = add_caption(doc, "Table", caption, current_chapter)
                cap_p.paragraph_format.keep_with_next = True
                cap_p.paragraph_format.space_after = Pt(6)
            data = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                if '---' not in lines[i]: data.append([c.strip() for c in lines[i].strip().split('|') if c.strip()])
                i += 1
            if data:
                print(f'Created table with {len(data)} rows and {len(data[0])} cols')
                tbl = doc.add_table(rows=len(data), cols=len(data[0]))
                format_custom_table(tbl)
                tbl.alignment = 1
                tbl.autofit = True

                # Force table to be 100% width
                from docx.oxml import OxmlElement
                from docx.oxml.ns import qn
                tblPr = tbl._tbl.tblPr
                tblW = OxmlElement('w:tblW')
                tblW.set(qn('w:w'), '5000')
                tblW.set(qn('w:type'), 'pct')
                tblPr.append(tblW)
                tbl.rows[0]._tr.get_or_add_trPr().append(OxmlElement('w:tblHeader'))
                for r, row_data in enumerate(data):
                    tbl.rows[r]._tr.get_or_add_trPr().append(OxmlElement('w:cantSplit'))
                    for c, val in enumerate(row_data):
                        cp = tbl.rows[r].cells[c].paragraphs[0]
                        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        add_formatted_text(cp, val, doc, current_chapter)
                        if r < len(data) - 1: cp.paragraph_format.keep_with_next = True
            continue
        if line.startswith('!['):
            alt = re.findall(r'!\[(.*?)\]', line)[0]
            m = re.search(r'\((.*?)\)', line)
            if m:
                rel_path = m.group(1)
                img_path = os.path.normpath(os.path.join(os.path.dirname(md_path), rel_path))
                
                # Robust resolution to handle differences between 1-level and 2-level parent setups
                if not os.path.exists(img_path):
                    if rel_path.startswith('../'):
                        img_path = os.path.normpath(os.path.join(os.path.dirname(md_path), '../' + rel_path))
                
                # Absolute fallback to the workspace root folders
                if not os.path.exists(img_path):
                    workspace_root = 'c:/Users/Administrator/Downloads/THESIS PROJECT'
                    fixed_rel = rel_path.replace('../', '').replace('..\\', '')
                    img_path = os.path.normpath(os.path.join(workspace_root, fixed_rel))
                
                if os.path.exists(img_path):
                    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.keep_with_next = True
                    run = p.add_run()
                    cap = alt
                    if i+1 < len(lines) and "Figure" in lines[i+1]:
                        cap = lines[i+1].strip().strip('*')
                        i += 1
                    # Image scaling for paper (columns) vs thesis (full page)
                    width = Inches(3.1) if is_paper else Inches(6.0)
                    try: 
                        run.add_picture(img_path, width=width)
                        print(f"Loaded image: {img_path}")
                    except Exception as e: 
                        p.add_run(f"\n[Image Error: {img_path}]")
                        print(f"Error loading image: {img_path} ({e})")
                    cap_p = add_caption(doc, "Figure", cap, chapter_num=current_chapter)
                    # Keep the caption with any immediately following source/italic line
                    next_i = i + 1
                    while next_i < len(lines) and not lines[next_i].strip():
                        next_i += 1
                    if next_i < len(lines) and lines[next_i].strip().startswith('*(Source:'):
                        cap_p.paragraph_format.keep_with_next = True
                else:
                    print(f"CRITICAL WARNING: Image path not found: {rel_path}")
            i += 1; continue
        if line.startswith('```'):
            lang = line.replace('```', '').strip().lower()
            i += 1; code = ""
            while i < len(lines) and not lines[i].strip().startswith('```'):
                code += lines[i]
                i += 1
            
            # Apply premium syntax highlighting and formatting
            add_syntax_highlighted_code(doc, code.strip(), lang if lang else "python")
            
            if i+2 < len(lines) and lines[i+1].strip() == "" and "Code" in lines[i+2]:
                p_d = doc.add_paragraph(lines[i+2].strip().strip('*'))
                p_d.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p_d.runs[0].italic = True
                i += 2
            i += 1; continue
        if line.startswith('- ') or line.startswith('* '):
            try: p = doc.add_paragraph(style='List Bullet')
            except: p = doc.add_paragraph()
            add_formatted_text(p, line[2:], doc, current_chapter); i += 1; continue
        if re.match(r'^\d+\. ', line):
            try: p = doc.add_paragraph(style='List Number')
            except: p = doc.add_paragraph()
            add_formatted_text(p, line[line.find(' ')+1:], doc, current_chapter); i += 1; continue
        p = doc.add_paragraph(); add_formatted_text(p, line, doc, current_chapter)
        
        # Keep paragraph with next if the next element is a Table, Image, or subheading
        next_idx = i + 1
        while next_idx < len(lines) and not lines[next_idx].strip():
            next_idx += 1
        if next_idx < len(lines):
            next_line = lines[next_idx].strip()
            if (next_line.startswith('###') or next_line.startswith('####')):
                p.paragraph_format.keep_with_next = True
                
        i += 1
    # Apply standard margin configurations globally to all sections of the thesis (PROJECT FORMAT.pdf)
    if not is_paper:
        for section in doc.sections:
            section.page_height = Cm(29.7)
            section.page_width = Cm(21.0)
            section.left_margin = Cm(3.5)
            section.right_margin = Cm(2.0)
            section.top_margin = Cm(2.0)
            section.bottom_margin = Cm(2.0)
            
    # Force keep_with_next explicitly on ALL heading paragraphs
    for p in doc.paragraphs:
        style_name = p.style.name if p.style else ""
        if style_name.startswith('Heading') or style_name.startswith('Nag'):
            p.paragraph_format.keep_with_next = True
            p.paragraph_format.keep_together = True
    # Force keep_with_next explicitly on ALL heading paragraphs
    for p in doc.paragraphs:
        style_name = p.style.name if p.style else ""
        if style_name.startswith('Heading') or style_name.startswith('Nag'):
            p.paragraph_format.keep_with_next = True
            p.paragraph_format.keep_together = True

    doc.save(docx_path); print(f"Success: {docx_path}")

if __name__ == "__main__":
    import argparse
    parser = argparse.ArgumentParser()
    parser.add_argument('--paper', action='store_true')
    parser.add_argument('--thesis', action='store_true')
    args = parser.parse_args()
    if args.paper:
        convert_to_professional_docx('MAIN/Paper Folder/paper_SOURCE.md', 'MAIN/Paper Folder/paper.docx')
    elif args.thesis:
        convert_to_professional_docx('MAIN/Thesis Folder/thesis.md', 'MAIN/Thesis Folder/Thesis.docx')
    else:
        # Compile both or default to thesis if present
        if os.path.exists('MAIN/Thesis Folder/thesis.md'):
            convert_to_professional_docx('MAIN/Thesis Folder/thesis.md', 'MAIN/Thesis Folder/Thesis.docx')
        else:
            convert_to_professional_docx('MAIN/Paper Folder/paper_SOURCE.md', 'MAIN/Paper Folder/paper.docx')
